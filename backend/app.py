from flask import Flask, render_template, request, redirect, url_for, session
from flask_cors import CORS
from werkzeug.utils import secure_filename
from datetime import datetime
import os
from dotenv import load_dotenv

from db import conn, cursor
from hash_utils import generate_hash
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document
from flask import send_from_directory
from user_agents import parse

from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from apscheduler.schedulers.background import BackgroundScheduler
from flask_mail import Mail, Message
from flask_socketio import SocketIO
import bcrypt
import threading


class MonitorHandler(FileSystemEventHandler):

    def on_modified(self, event):

        if not event.is_directory:

            print("Modified:", event.src_path)

            add_log(
                f"Real-Time Modification Detected: {event.src_path}"
            )
        
load_dotenv()
app = Flask(

    __name__,

    template_folder="../frontend",

    static_folder="../frontend"

)

CORS(app)

app.secret_key = os.getenv("SECRET_KEY")

UPLOAD_FOLDER = os.path.join(
    os.path.dirname(__file__),
    "../uploads"
)

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# EMAIL ALERT CONFIGURATION

app.config['MAIL_SERVER'] = 'smtp.gmail.com'

app.config['MAIL_PORT'] = 587

app.config['MAIL_USE_TLS'] = True

app.config['MAIL_USERNAME'] = os.getenv("MAIL_USERNAME")

app.config['MAIL_PASSWORD'] = os.getenv("MAIL_PASSWORD")
mail = Mail(app)
socketio = SocketIO(app)

# ======================================================
# SEND MAIL IN BACKGROUND
# ======================================================

def send_async_email(app, msg):

    with app.app_context():

        mail.send(msg)

# ======================================================
# REAL TIME ALERT EMITTER
# ======================================================

def emit_realtime_alert(message):

    socketio.emit(

        'new_alert',

        {

            'message': message

        }

    )

# REAL-TIME FILE MONITORING

observer = Observer()

observer.schedule(

    MonitorHandler(),

    path=UPLOAD_FOLDER,

    recursive=True

)

observer.start()


from datetime import datetime

def current_time():

    return datetime.now().strftime(
        "%Y-%m-%d %H:%M:%S"
    )

# ======================================================
# LOG FUNCTION
# ======================================================

def add_log(activity):

    cursor.execute(
        """
        INSERT INTO logs(
            activity,
            ip_address,
            created_at
        )
        VALUES(%s,%s,%s)
        """,
        (
            activity,
            request.remote_addr,
            datetime.now()
        )
    )

    conn.commit()

    


# ======================================================
# HOME PAGE
# ======================================================

@app.route('/')
def home():

    return redirect(url_for('login'))

# ======================================================
# LOGIN MODULE
# ======================================================

@app.route('/login', methods=['GET', 'POST'])
def login():

    # CHECK BLOCKED IP

    if is_ip_blocked(request.remote_addr):

        add_log(
            f"Blocked IP Tried Access: {request.remote_addr}"
        )

        return render_template(
            "access_denied.html"
        )

    if request.method == 'POST':

        username = request.form['username']

        password = request.form['password']

        # GET ADMIN BY USERNAME

        cursor.execute(
            """
            SELECT * FROM admin_users
            WHERE username=%s
            """,
            (username,)
        )

        admin = cursor.fetchone()

        # SUCCESSFUL LOGIN

        if admin and bcrypt.checkpw(

            password.encode('utf-8'),

            admin['password'].encode('utf-8')

        ):

            session['admin_logged_in'] = True

            session['admin_name'] = admin['username']

            session['role'] = admin['role']

            # STORE LOGIN DEVICE + IP

            user_agent = parse(
                request.headers.get('User-Agent')
            )

            browser = user_agent.browser.family

            os_name = user_agent.os.family

            device = user_agent.device.family

            device_info = f"{os_name} | {browser} | {device}"

            cursor.execute(
                """
                UPDATE admin_users
                SET
                last_login_ip=%s,
                device_info=%s,
                last_login=%s
                WHERE id=%s
                """,
                (
                    request.remote_addr,
                    device_info,
                    datetime.now(),
                    admin['id']
                )
            )

            conn.commit()

            add_log(
                f"{admin['username']} Logged In"
            )

            return redirect(
                url_for('dashboard')
            )

        # FAILED LOGIN

        else:

            add_log(
                f"Failed Login Attempt for username: {username}"
            )

            cursor.execute(
                """
                INSERT INTO alerts(
                    attack_type,
                    attacker_ip,
                    severity,
                    message,
                    created_at
                )
                VALUES(%s,%s,%s,%s,%s)
                """,
                (
                    "Unauthorized Login",

                    request.remote_addr,

                    "High",

                    f"Failed login attempt detected for username: {username}",

                    datetime.now()
                )
            )

            # SUSPICIOUS IP TRACKING

            cursor.execute(
                """
                SELECT * FROM suspicious_ips
                WHERE ip_address=%s
                """,
                (request.remote_addr,)
            )

            existing_ip = cursor.fetchone()

            # BRUTE FORCE DETECTION

            if existing_ip and existing_ip['failed_attempts'] >= 5:

                cursor.execute(
                    """
                    INSERT INTO alerts(
                        attack_type,
                        attacker_ip,
                        severity,
                        message,
                        created_at
                    )
                    VALUES(%s,%s,%s,%s,%s)
                    """,
                    (
                        "Brute Force Attack",

                        request.remote_addr,

                        "Critical",

                        f"5+ failed login attempts detected from IP: {request.remote_addr}",

                        datetime.now()
                    )
                )

                add_log(
                    f"Brute Force Attack Detected From {request.remote_addr}"
                )

                emit_realtime_alert(
                    "Brute Force Attack Detected"
                )

            # UPDATE FAILED ATTEMPTS

            if existing_ip:

                cursor.execute(
                    """
                    UPDATE suspicious_ips
                    SET failed_attempts=failed_attempts+1
                    WHERE ip_address=%s
                    """,
                    (request.remote_addr,)
                )

            else:

                cursor.execute(
                    """
                    INSERT INTO suspicious_ips(
                        ip_address
                    )
                    VALUES(%s)
                    """,
                    (request.remote_addr,)
                )

            conn.commit()

            # REALTIME ALERT

            emit_realtime_alert(
                "Unauthorized Login Attempt"
            )

            return render_template(
                'login.html',
                error="Invalid Username or Password"
            )

    return render_template('login.html')
# ======================================================
# DASHBOARD
# ======================================================

@app.route('/dashboard')
def dashboard():

    if not session.get('admin_logged_in'):

        return redirect(url_for('login'))

    cursor.execute(
        """
        SELECT * FROM files
        ORDER BY id DESC
        """
    )

    files = cursor.fetchall()

    cursor.execute(
        """
        SELECT * FROM alerts
        ORDER BY id DESC
        """
    )

    alerts = cursor.fetchall()

    secure_count = sum(
        1 for f in files
        if f['status'] == 'Secure'
    )

    modified_count = sum(
        1 for f in files
        if f['status'] == 'Modified'
    )

    missing_count = sum(
        1 for f in files
        if f['status'] == 'Missing'
    )

    # OPEN ALERT COUNT

    cursor.execute(
        """
        SELECT COUNT(*) AS total
        FROM alerts
        WHERE status='Open'
        """
    )

    alert_count = cursor.fetchone()['total']

    return render_template(

        'dashboard.html',

        request=request,

        files=files,

        alerts=alerts,

        total_files=len(files),

        secure_count=secure_count,

        modified_count=modified_count,

        missing_count=missing_count,

        alert_count=alert_count
    )


@app.route('/files/<status>')
def filtered_files(status):

    if status == "all":

        cursor.execute("""
        SELECT * FROM files
        ORDER BY id DESC
        """)

    else:

        cursor.execute("""
        SELECT * FROM files
        WHERE status=%s
        ORDER BY id DESC
        """, (status,))

    files = cursor.fetchall()

    return render_template(
        'filtered_files.html',
        files=files,
        status=status
    )


# ======================================================
# FILE UPLOAD
# ======================================================

@app.route('/upload', methods=['POST'])
def upload_file():

    file = request.files['file']

    if file.filename == '':

        return "No File Selected"

    filename = secure_filename(file.filename)

    # CHECK DUPLICATE FILE

    cursor.execute(
        """
        SELECT * FROM files
        WHERE file_name=%s
        """,
        (
            filename,
        )
    )

    existing_file = cursor.fetchone()

    if existing_file:

        return "File Already Exists In Monitoring System"

    file_path = os.path.join(
        app.config['UPLOAD_FOLDER'],
        filename
    )

    file.save(file_path)

    hash_value = generate_hash(file_path)

    cursor.execute(
        """
        INSERT INTO files(
            file_name,
            file_path,
            hash_value,
            status,
            uploaded_at
        )
        VALUES(%s,%s,%s,%s,%s)
        """,
        (
            filename,
            file_path,
            hash_value,
            'Secure',
            datetime.now()
        )
    )

    conn.commit()

    add_log(f"{filename} uploaded")

    return redirect(url_for('dashboard'))

from flask import send_from_directory
@app.route('/uploads/<filename>')
def uploaded_file(filename):

    return send_from_directory(
        app.config['UPLOAD_FOLDER'],
        filename
    )


# ======================================================
# FILE MONITORING
# ======================================================

@app.route('/monitor')
def monitor_files():

    cursor.execute("""
    SELECT * FROM files
    ORDER BY id DESC
    """)

    files = cursor.fetchall()

    for file in files:

        file_path = file['file_path']

        old_hash = file['hash_value']

        # ==================================================
        # FILE DELETION CHECK
        # ==================================================

        if not os.path.exists(file_path):

            cursor.execute(
                """
                UPDATE files
                SET status=%s
                WHERE id=%s
                """,
                (
                    'Missing',
                    file['id']
                )
            )

            # PREVENT DUPLICATE DELETE ALERTS

            cursor.execute(
                """
                SELECT * FROM alerts
                WHERE file_id=%s
                AND attack_type='File Deletion'
                ORDER BY id DESC
                LIMIT 1
                """,
                (
                    file['id'],
                )
            )

            existing_delete_alert = cursor.fetchone()

            if not existing_delete_alert:

                cursor.execute(
                    """
                    INSERT INTO alerts(
                        file_id,
                        attack_type,
                        attacker_ip,
                        severity,
                        message,
                        created_at
                    )
                    VALUES(%s,%s,%s,%s,%s,%s)
                    """,
                    (
                        file['id'],
                        'File Deletion',
                        request.remote_addr,
                        'Critical',
                        f"{file['file_name']} was deleted",
                        datetime.now()
                    )
                )

                add_log(
                    f"{file['file_name']} deleted"
                )

                conn.commit()

            continue

        # ==================================================
        # HASH CHECK
        # ==================================================

        current_hash = generate_hash(file_path)

        cursor.execute(
            """
            SELECT * FROM alerts
            WHERE file_id=%s
            AND attack_type='File Tampering'
            ORDER BY id DESC
            LIMIT 1
            """,
            (
                file['id'],
            )
        )

        existing_alert = cursor.fetchone()

        # ==================================================
        # FILE MODIFIED
        # ==================================================

        if current_hash != old_hash:

            # UPDATE STATUS + STORE NEW HASH

            cursor.execute(
                """
                UPDATE files
                SET
                status=%s,
                current_hash=%s
                WHERE id=%s
                """,
                (
                    'Modified',
                    current_hash,
                    file['id']
                )
            )

            # INSERT ALERT ONLY ONCE

            if not existing_alert:

                cursor.execute(
                    """
                    INSERT INTO alerts(
                        file_id,
                        attack_type,
                        attacker_ip,
                        severity,
                        message,
                        created_at
                    )
                    VALUES(%s,%s,%s,%s,%s,%s)
                    """,
                    (
                        file['id'],
                        'File Tampering',
                        request.remote_addr,
                        'High',
                        f"{file['file_name']} was modified",
                        datetime.now()
                    )
                )

                add_log(
                    f"{file['file_name']} modified"
                )

            conn.commit()

        # ==================================================
        # FILE SECURE
        # ==================================================

        else:

            cursor.execute(
                """
                UPDATE files
                SET status=%s
                WHERE id=%s
                """,
                (
                    'Secure',
                    file['id']
                )
            )

            conn.commit()

    return redirect(url_for('dashboard'))

# ======================================================
# ALERT PAGE
# ======================================================

@app.route('/alerts')
def alerts_page():

    cursor.execute(
        """
        SELECT * FROM alerts
        ORDER BY id DESC
        """
    )

    alerts = cursor.fetchall()

    return render_template(
        'alerts.html',
        alerts=alerts
    )


@app.route('/resolve-alert/<int:id>')
def resolve_alert(id):

    cursor.execute(
        """
        UPDATE alerts
        SET status='Resolved'
        WHERE id=%s
        """,
        (id,)
    )

    conn.commit()

    add_log(
        f"Resolved Alert ID {id}"
    )

    return redirect(
        url_for('alerts_page')
    )

# ======================================================
# REPORT GENERATOR
# ======================================================

@app.route('/generate-report')
def generate_report():

    cursor.execute(
        """
        SELECT * FROM alerts
        """
    )

    alerts = cursor.fetchall()

    for alert in alerts:

        details = f"""
Attack Type: {alert['attack_type']}
Severity: {alert['severity']}
IP Address: {alert['attacker_ip']}
Message: {alert['message']}
Timestamp: {alert['created_at']}
"""

        cursor.execute(
            """
            INSERT INTO reports(
                alert_id,
                status,
                details
            )
            VALUES(%s,%s,%s)
            """,
            (
                alert['id'],
                'Generated',
                details
            )
        )

    conn.commit()

    add_log("Security Report Generated")

    return redirect(url_for('reports_page'))


@app.route('/download-report-pdf')
def download_report_pdf():

    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet

    cursor.execute(
        """
        SELECT * FROM reports
        ORDER BY id DESC
        """
    )

    reports = cursor.fetchall()

    pdf_path = "security_report.pdf"

    doc = SimpleDocTemplate(pdf_path)

    styles = getSampleStyleSheet()

    elements = []

    title = Paragraph(
        "FIM Security Incident Report",
        styles['Title']
    )

    elements.append(title)

    for report in reports:

        text = f"""
        Status: {report['status']}<br/>
        Details: {report['details']}<br/><br/>
        """

        elements.append(
            Paragraph(
                text,
                styles['BodyText']
            )
        )

    doc.build(elements)

    add_log("PDF Report Downloaded")

    from flask import send_file

    return send_file(
        pdf_path,
        as_attachment=True
    )


# ======================================================
# REPORT PAGE
# ======================================================

@app.route('/reports')
def reports_page():

    cursor.execute(
        """
        SELECT * FROM reports
        ORDER BY id DESC
        """
    )

    reports = cursor.fetchall()

    return render_template(
        'reports.html',
        reports=reports
    )

@app.route('/download-report-doc')
def download_report_doc():

    document = Document()

    document.add_heading(
        'FIM Security Incident Report',
        0
    )

    cursor.execute(
        """
        SELECT * FROM reports
        ORDER BY id DESC
        """
    )

    reports = cursor.fetchall()

    for report in reports:

        document.add_paragraph(

            f"""
Status: {report['status']}

Details:
{report['details']}
            """
        )

    doc_path = "security_report.docx"

    document.save(doc_path)

    add_log("DOC Report Downloaded")

    from flask import send_file

    return send_file(
        doc_path,
        as_attachment=True
    )



@app.route('/generate-single-report/<int:alert_id>')
def generate_single_report(alert_id):

    cursor.execute(
        """
        SELECT * FROM alerts
        WHERE id=%s
        """,
        (
            alert_id,
        )
    )

    alert = cursor.fetchone()

    if not alert:

        return "Alert Not Found"

    details = f"""
Attack Type: {alert['attack_type']}

Severity: {alert['severity']}

IP Address: {alert['attacker_ip']}

Message: {alert['message']}

Timestamp: {alert['created_at']}
"""

    cursor.execute(
        """
        INSERT INTO reports(
            alert_id,
            status,
            details
        )
        VALUES(%s,%s,%s)
        """,
        (
            alert['id'],
            'Generated',
            details
        )
    )

    conn.commit()

    add_log(
        f"Generated Report For Alert ID {alert_id}"
    )

    return redirect(url_for('reports_page'))

# ======================================================
# BLOCKED IPS PAGE
# ======================================================

@app.route('/blocked-ips')
def blocked_ips_page():

    cursor.execute(
        """
        SELECT * FROM blocked_ips
        ORDER BY id DESC
        """
    )

    ips = cursor.fetchall()

    return render_template(
        'blocked_ips.html',
        ips=ips
    )


def is_ip_blocked(ip):

    cursor.execute(
        """
        SELECT * FROM blocked_ips
        WHERE ip_address=%s
        """,
        (ip,)
    )

    return cursor.fetchone()


@app.route("/block-ip", methods=["POST"])
def block_ip():

    ip = request.form["ip"]

    # CHECK ALREADY BLOCKED

    cursor.execute(
        """
        SELECT * FROM blocked_ips
        WHERE ip_address=%s
        """,
        (
            ip,
        )
    )

    existing_ip = cursor.fetchone()

    # INSERT ONLY IF NOT EXISTS

    if not existing_ip:

        cursor.execute(
            """
            INSERT INTO blocked_ips
            (
                ip_address,
                reason,
                blocked_at
            )
            VALUES (%s,%s,%s)
            """,
            (
                ip,
                "Verified Threat",
                current_time()
            )
        )

        conn.commit()

        add_log(f"Blocked IP {ip}")

    return redirect("/blocked-ips")



@app.route('/unblock-ip/<int:id>')
def unblock_ip(id):

    if session.get('role') != 'SuperAdmin':
        return "Access Denied"

    cursor.execute("""
        SELECT *
        FROM blocked_ips
        WHERE id=%s
    """, (id,))

    blocked_ip = cursor.fetchone()

    if not blocked_ip:
        return redirect('/blocked-ips')

    cursor.execute("""
        DELETE FROM blocked_ips
        WHERE id=%s
    """, (id,))

    conn.commit()

    add_log(
        f"IP Unblocked: {blocked_ip['ip_address']}"
    )

    return redirect('/blocked-ips')


# ======================================================
# LOGS PAGE
# ======================================================

@app.route('/logs')
def logs_page():

    cursor.execute(
        """
        SELECT * FROM logs
        ORDER BY id DESC
        """
    )

    logs = cursor.fetchall()

    return render_template(
        'logs.html',
        logs=logs
    )



@app.route('/clear-alerts')
def clear_alerts():

    # ONLY SUPERADMIN

    if session.get('role') != 'SuperAdmin':

        return "Access Denied"

    cursor.execute(
        "DELETE FROM alerts"
    )

    conn.commit()

    add_log(
        "Alerts History Cleared By SuperAdmin"
    )

    return redirect(
        url_for('alerts_page')
    )

@app.route('/clear-logs')
def clear_logs():

    # ONLY SUPERADMIN

    if session.get('role') != 'SuperAdmin':

        return "Access Denied"

    cursor.execute("DELETE FROM logs")

    conn.commit()

    add_log("Logs Cleared By SuperAdmin")

    return redirect(url_for('logs_page'))


@app.route('/clear-reports')
def clear_reports():

    cursor.execute("DELETE FROM reports")

    conn.commit()

    add_log("Reports Cleared")

    return redirect(url_for('reports_page'))


@app.route('/clear-blocked-ips')
def clear_blocked_ips():

    # ONLY SUPERADMIN

    if session.get('role') != 'SuperAdmin':

        return "Access Denied"

    cursor.execute("DELETE FROM blocked_ips")

    conn.commit()

    add_log("Blocked IP History Cleared")

    return redirect(url_for('blocked_ips_page'))

# ======================================================
# USER MANAGEMENT
# ======================================================

@app.route('/admin-management')
def admin_management():

    # ONLY SUPERADMIN

    if session.get('role') != 'SuperAdmin':

        return redirect(url_for('dashboard'))

    cursor.execute(
        """
        SELECT * FROM admin_users
        ORDER BY id DESC
        """
    )

    admins = cursor.fetchall()

    return render_template(
        'admin_management.html',
        admins=admins
    )


# ======================================================
# ADD ADMIN
# ======================================================

@app.route('/add-admin', methods=['POST'])
def add_admin():

    # ONLY SUPERADMIN

    if session.get('role') != 'SuperAdmin':

        return "Access Denied"

    username = request.form['username']

    password = request.form['password']
    role = request.form['role']

    hashed_password = bcrypt.hashpw(

    password.encode('utf-8'),

    bcrypt.gensalt()

).decode('utf-8')

    # CHECK EXISTING USER

    cursor.execute(
        """
        SELECT * FROM admin_users
        WHERE username=%s
        """,
        (
            username,
        )
    )

    existing_admin = cursor.fetchone()

    if existing_admin:

        return "Admin Already Exists"

    cursor.execute(
        """
        INSERT INTO admin_users(
            username,
            password,
            role
        )
        VALUES(%s,%s,%s)
        """,
        (
            username,
            hashed_password,
            role
        )
    )

    conn.commit()

    add_log(
        f"New Admin Added: {username}"
    )

    return redirect(
        url_for('admin_management')
    )


# ======================================================
## DELETE ADMIN
# ======================================================

@app.route('/delete-admin/<int:id>')
def delete_admin(id):

    # ONLY SUPERADMIN

    if session.get('role') != 'SuperAdmin':

        return "Access Denied"

    # GET ADMIN DETAILS

    cursor.execute(
        """
        SELECT * FROM admin_users
        WHERE id=%s
        """,
        (
            id,
        )
    )

    admin = cursor.fetchone()

    # CHECK ADMIN EXISTS

    if not admin:

        return "Admin Not Found"

    # PREVENT SELF DELETE

    if admin['username'] == session.get('admin_name'):

        return "You Cannot Delete Yourself"

    # CHECK TOTAL SUPERADMINS

    cursor.execute(
        """
        SELECT COUNT(*) AS total
        FROM admin_users
        WHERE role='SuperAdmin'
        """
    )

    total_superadmins = cursor.fetchone()['total']

    # PREVENT LAST SUPERADMIN DELETE

    if admin['role'] == 'SuperAdmin' and total_superadmins <= 1:

        return "Cannot Delete Last SuperAdmin"

    # DELETE ADMIN

    cursor.execute(
        """
        DELETE FROM admin_users
        WHERE id=%s
        """,
        (
            id,
        )
    )

    conn.commit()

    add_log(
        f"Admin Deleted ID: {id}"
    )

    return redirect(
        url_for('admin_management')
    )


# ======================================================
# LOGOUT
# ======================================================

@app.route('/logout')
def logout():

    add_log("Admin Logged Out")

    session.clear()

    return redirect(url_for('login'))

@app.route('/hash-details/<int:file_id>')
def hash_details(file_id):

    cursor.execute("""
    SELECT * FROM files
    WHERE id=%s
    """, (file_id,))

    file = cursor.fetchone()

    return render_template(
        'hash_details.html',
        file=file
    )

@app.route('/test-mail')
def test_mail():

    msg = Message(

        'CyberShield Test',

        sender=os.getenv("MAIL_USERNAME"),

        recipients=[os.getenv("MAIL_USERNAME")]

    )

    msg.body = "CyberShield email alert working successfully."

    mail.send(msg)

    return "Email Sent Successfully"
    # ======================================================
# ======================================================
# SECURITY EVENT API
# ======================================================

@app.route('/api/security-alert', methods=['POST'])
def api_security_alert():

    # API KEY VALIDATION

    api_key = request.headers.get('X-API-KEY')

    if api_key != os.getenv("API_KEY"):

        return {
            "status": "error",
            "message": "Unauthorized API Access"
        }, 401

    data = request.json

    attack_type = data.get('attack_type')

    attacker_ip = data.get('ip')

    severity = data.get('severity')

    message = data.get('message')

    source_app = data.get('source_app')

    # INSERT ALERT

    cursor.execute(
        """
        INSERT INTO alerts(
            attack_type,
            attacker_ip,
            severity,
            message,
            created_at
        )
        VALUES(%s,%s,%s,%s,%s)
        """,
        (
            attack_type,
            attacker_ip,
            severity,
            f"[{source_app}] {message}",
            datetime.now()
        )
    )

    conn.commit()

    # REAL TIME ALERT

    emit_realtime_alert(
        f"{source_app} : {attack_type}"
    )

    # LOG

    add_log(
        f"API Alert Received From {source_app}"
    )

    # EMAIL ALERT

    try:

        msg = Message(

            'CyberShield API Threat Alert',

            sender=app.config['MAIL_USERNAME'],

            recipients=[os.getenv("MAIL_USERNAME")]

        )

        msg.body = f"""

CyberShield API Threat Alert

Source Application: {source_app}

Attack Type: {attack_type}

IP Address: {attacker_ip}

Severity: {severity}

Message: {message}

Time: {datetime.now()}

"""

        # SEND MAIL IN BACKGROUND

        threading.Thread(

            target=send_async_email,

            args=(app, msg)

        ).start()

    except Exception as e:

        print("MAIL ERROR:", e)

    return {
        "status": "success",
        "message": "Alert Received"
    }


@app.route('/profile')
def profile():

    if not session.get('admin_logged_in'):
        return redirect(url_for('login'))

    cursor.execute("""
        SELECT *
        FROM admin_users
        WHERE username=%s
    """, (session['admin_name'],))

    admin = cursor.fetchone()

    return render_template(
        'profile.html',
        admin=admin
    )


@app.route('/change-password', methods=['GET', 'POST'])
def change_password():

    if not session.get('admin_logged_in'):
        return redirect(url_for('login'))

    if request.method == 'POST':

        current_password = request.form['current_password']
        new_password = request.form['new_password']
        confirm_password = request.form['confirm_password']

        if new_password != confirm_password:
            return render_template(
                'change_password.html',
                error="Passwords do not match"
            )

        cursor.execute("""
            SELECT *
            FROM admin_users
            WHERE username=%s
        """, (session['admin_name'],))

        admin = cursor.fetchone()

        if not bcrypt.checkpw(
            current_password.encode('utf-8'),
            admin['password'].encode('utf-8')
        ):
            return render_template(
                'change_password.html',
                error="Current Password Incorrect"
            )

        hashed_password = bcrypt.hashpw(
            new_password.encode('utf-8'),
            bcrypt.gensalt()
        ).decode('utf-8')

        cursor.execute("""
            UPDATE admin_users
            SET password=%s
            WHERE id=%s
        """, (
            hashed_password,
            admin['id']
        ))

        conn.commit()

        add_log(
            f"{admin['username']} Changed Password"
        )

        return render_template(
            'change_password.html',
            success="Password Changed Successfully"
        )

    return render_template('change_password.html')


@app.route('/reset-password/<int:id>')
def reset_password(id):

    if session.get('role') != 'SuperAdmin':
        return "Access Denied"

    new_password = "admin@123"

    hashed_password = bcrypt.hashpw(
        new_password.encode('utf-8'),
        bcrypt.gensalt()
    ).decode('utf-8')

    cursor.execute("""
        UPDATE admin_users
        SET password=%s
        WHERE id=%s
    """, (
        hashed_password,
        id
    ))

    conn.commit()

    add_log(
        f"Password Reset For Admin ID {id}"
    )

    return redirect('/admin-management')

    
# ======================================================
# MAIN
# ======================================================

if __name__ == '__main__':

    socketio.run(

        app,

        host="0.0.0.0",

        port=5000,

        debug=False

    )